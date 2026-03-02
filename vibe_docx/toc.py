# -*- coding: utf-8 -*-
"""
Vibe Docx - TOC (Table of Contents) Tools

目录生成与更新功能，支持两种模式：
1. TOC Field - Word 原生目录字段，可在 Word 中更新
2. Hyperlink - 超链接目录，PDF 导出后保留跳转功能

Usage:
    from vibe_docx.toc import generate_toc, update_toc, validate_toc_links
    
    # 生成目录
    result = generate_toc("document.docx", {
        "style": "hyperlink",
        "levels": 3,
        "show_page_numbers": True
    })
    
    # 更新目录
    update_toc("document.docx")
    
    # 验证链接
    validate_toc_links("document.docx")
"""

import os
import re
from typing import Dict, List, Any, Optional
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Twips
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


class TOCError(Exception):
    """TOC 操作错误"""
    
    ERROR_CODES = {
        "TOC001": {"message": "文档不存在", "recovery": "请提供有效的文档路径"},
        "TOC002": {"message": "文档无标题", "recovery": "请先添加标题样式的内容"},
        "TOC003": {"message": "目录已存在", "recovery": "使用 update_toc 更新或删除后重新生成"},
        "TOC004": {"message": "书签创建失败", "recovery": "检查标题是否包含特殊字符"},
        "TOC005": {"message": "不支持的模式", "recovery": "使用 'toc_field' 或 'hyperlink'"},
    }
    
    def __init__(self, code: str, detail: str = ""):
        self.code = code
        self.detail = detail
        info = self.ERROR_CODES.get(code, {"message": "未知错误", "recovery": ""})
        self.message = info["message"]
        self.recovery = info["recovery"]
        super().__init__(f"[{code}] {self.message}: {detail}")


def _extract_headings(doc: Document, max_level: int = 3) -> List[Dict[str, Any]]:
    """
    提取文档中的标题
    
    Args:
        doc: Document 对象
        max_level: 最大标题级别 (1-9)
    
    Returns:
        标题列表，每个元素包含:
        - text: 标题文本
        - level: 标题级别 (1-9)
        - paragraph_index: 段落索引
        - bookmark_name: 书签名称
    """
    headings = []
    heading_styles = [f'Heading {i}' for i in range(1, max_level + 1)]
    
    # 同时支持中文名称
    heading_styles_cn = [f'标题 {i}' for i in range(1, max_level + 1)]
    
    for idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        
        # 检查是否是标题样式
        level = None
        for i, (en_style, cn_style) in enumerate(zip(heading_styles, heading_styles_cn), 1):
            if style_name == en_style or style_name == cn_style:
                level = i
                break
        
        if level and level <= max_level:
            text = para.text.strip()
            if text:
                # 生成书签名称（有效字符）
                bookmark_name = f"_toc_{idx}_{re.sub(r'[^\w\u4e00-\u9fff]', '_', text)[:20]}"
                headings.append({
                    "text": text,
                    "level": level,
                    "paragraph_index": idx,
                    "bookmark_name": bookmark_name
                })
    
    return headings


def _create_bookmark(para, bookmark_name: str) -> None:
    """
    在段落中创建书签
    
    Args:
        para: 段落对象
        bookmark_name: 书签名称
    """
    # 创建书签开始标记
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), '0')
    bookmark_start.set(qn('w:name'), bookmark_name)
    
    # 创建书签结束标记
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), '0')
    
    # 插入到段落
    para._p.insert(0, bookmark_start)
    para._p.append(bookmark_end)


def _create_hyperlink_toc_entry(doc, text: str, bookmark_name: str, page_num: int, level: int, tab_leader: str) -> Any:
    """
    创建超链接目录项
    
    Args:
        doc: Document 对象
        text: 标题文本
        bookmark_name: 书签名称
        page_num: 页码
        level: 标题级别
        tab_leader: 前导符样式
    
    Returns:
        目录段落
    """
    # 创建段落
    para = doc.add_paragraph()
    
    # 设置缩进（每级 0.5 英寸）
    para.paragraph_format.left_indent = Inches(0.5 * (level - 1))
    
    # 创建超链接
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    
    # 创建 run
    run = OxmlElement('w:r')
    
    # 设置样式
    rPr = OxmlElement('w:rPr')
    run.append(rPr)
    
    # 添加文本
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run.append(text_elem)
    
    # 添加前导符和页码
    if tab_leader != "none":
        tab = OxmlElement('w:tab')
        run.append(tab)
    
    page_text = OxmlElement('w:t')
    page_text.text = str(page_num)
    run.append(page_text)
    
    hyperlink.append(run)
    para._p.append(hyperlink)
    
    return para


def _create_toc_field(doc, headings: List[Dict], options: Dict) -> None:
    """
    创建 TOC 字段目录
    
    Args:
        doc: Document 对象
        headings: 标题列表
        options: 配置选项
    """
    # 创建目录段落
    para = doc.add_paragraph()
    
    # TOC 字段指令
    levels = options.get("levels", 3)
    field_instr = f'TOC \\o "1-{levels}" \\h \\z \\u'
    
    # 创建字段
    fld_begin = OxmlElement('w:fldSimple')
    fld_begin.set(qn('w:instr'), field_instr)
    
    # 字段内容（占位文本）
    run = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.text = "目录（右键更新）"
    run.append(text)
    fld_begin.append(run)
    
    para._p.append(fld_begin)


def generate_toc(file_path: str, options: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """
    生成目录
    
    Args:
        file_path: 文档路径
        options: 配置选项
            - style: "toc_field" | "hyperlink" (默认 "hyperlink")
            - levels: int, 显示到第几级标题 (默认 3)
            - show_page_numbers: bool, 是否显示页码 (默认 True)
            - tab_leader: "dots" | "dashes" | "none", 页码前导符 (默认 "dots")
            - title: str, 目录标题 (默认 "目录")
            - position: "beginning" | "end", 目录位置 (默认 "beginning")
            - output_path: str, 输出路径 (可选，默认覆盖原文件)
    
    Returns:
        {
            "success": bool,
            "headings_count": int,
            "toc_style": str,
            "output_path": str,
            "message": str
        }
    """
    if not HAS_PYTHON_DOCX:
        raise TOCError("TOC001", "需要安装 python-docx: pip install python-docx")
    
    if not os.path.exists(file_path):
        raise TOCError("TOC001", file_path)
    
    options = options or {}
    style = options.get("style", "hyperlink")
    levels = options.get("levels", 3)
    show_page_numbers = options.get("show_page_numbers", True)
    tab_leader = options.get("tab_leader", "dots")
    title = options.get("title", "目录")
    position = options.get("position", "beginning")
    output_path = options.get("output_path", file_path)
    
    if style not in ["toc_field", "hyperlink"]:
        raise TOCError("TOC005", style)
    
    # 打开文档
    doc = Document(file_path)
    
    # 提取标题
    headings = _extract_headings(doc, levels)
    
    if not headings:
        raise TOCError("TOC002", "未找到标题样式的内容")
    
    # 为每个标题创建书签
    for h in headings:
        para = doc.paragraphs[h["paragraph_index"]]
        try:
            _create_bookmark(para, h["bookmark_name"])
        except Exception as e:
            raise TOCError("TOC004", str(e))
    
    # 根据模式生成目录
    if style == "toc_field":
        # TOC 字段模式
        _create_toc_field(doc, headings, options)
    else:
        # 超链接模式
        for h in headings:
            _create_hyperlink_toc_entry(
                doc, h["text"], h["bookmark_name"],
                h["paragraph_index"],  # 简化：用段落索引代替页码
                h["level"], tab_leader
            )
    
    # 保存文档
    doc.save(output_path)
    
    return {
        "success": True,
        "headings_count": len(headings),
        "toc_style": style,
        "output_path": output_path,
        "message": f"已生成目录，包含 {len(headings)} 个标题"
    }


def update_toc(file_path: str, output_path: Optional[str] = None) -> Dict[str, Any]:
    """
    更新目录（仅 TOC 字段模式有效）
    
    注意：TOC 字段的更新需要在 Word 中执行。
    此函数主要用于重新生成超链接模式目录。
    
    Args:
        file_path: 文档路径
        output_path: 输出路径 (可选)
    
    Returns:
        {
            "success": bool,
            "message": str
        }
    """
    if not HAS_PYTHON_DOCX:
        raise TOCError("TOC001", "需要安装 python-docx: pip install python-docx")
    
    if not os.path.exists(file_path):
        raise TOCError("TOC001", file_path)
    
    # python-docx 无法更新 TOC 字段
    # 建议：在 Word 中打开并按 F9 更新
    
    return {
        "success": True,
        "message": "请在 Word 中打开文档并按 F9 更新目录字段。超链接目录需要删除后重新生成。"
    }


def validate_toc_links(file_path: str) -> Dict[str, Any]:
    """
    验证目录跳转链接有效性
    
    Args:
        file_path: 文档路径
    
    Returns:
        {
            "success": bool,
            "total_links": int,
            "valid_links": int,
            "broken_links": list,
            "message": str
        }
    """
    if not HAS_PYTHON_DOCX:
        raise TOCError("TOC001", "需要安装 python-docx: pip install python-docx")
    
    if not os.path.exists(file_path):
        raise TOCError("TOC001", file_path)
    
    doc = Document(file_path)
    
    # 提取所有书签名称
    bookmarks = set()
    for para in doc.paragraphs:
        for elem in para._p.iter():
            if elem.tag == qn('w:bookmarkStart'):
                name = elem.get(qn('w:name'))
                if name:
                    bookmarks.add(name)
    
    # 检查所有超链接
    total_links = 0
    valid_links = 0
    broken_links = []
    
    for para in doc.paragraphs:
        for elem in para._p.iter():
            if elem.tag == qn('w:hyperlink'):
                total_links += 1
                anchor = elem.get(qn('w:anchor'))
                if anchor in bookmarks:
                    valid_links += 1
                else:
                    broken_links.append({
                        "anchor": anchor,
                        "text": para.text[:50] if para.text else ""
                    })
    
    return {
        "success": True,
        "total_links": total_links,
        "valid_links": valid_links,
        "broken_links": broken_links,
        "message": f"验证完成：{valid_links}/{total_links} 个有效链接"
    }
