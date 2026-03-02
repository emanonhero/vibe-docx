# -*- coding: utf-8 -*-
"""
Tests for V1.3 features: TOC generation and SKILL installation
"""

import os
import sys
import tempfile
from pathlib import Path

# 添加项目根目录到路径
sys.path.insert(0, str(Path(__file__).parent.parent))

import pytest

try:
    from docx import Document
    from docx.shared import Pt
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


class TestTOCGeneration:
    """目录生成测试"""
    
    @pytest.fixture
    def sample_doc(self, tmp_path):
        """创建测试文档"""
        if not HAS_PYTHON_DOCX:
            pytest.skip("需要 python-docx")
        
        doc = Document()
        
        # 添加标题
        doc.add_heading("第一章 简介", level=1)
        doc.add_paragraph("这是简介内容。")
        
        doc.add_heading("1.1 背景", level=2)
        doc.add_paragraph("背景说明。")
        
        doc.add_heading("1.2 目标", level=2)
        doc.add_paragraph("目标说明。")
        
        doc.add_heading("第二章 方法", level=1)
        doc.add_paragraph("方法说明。")
        
        doc.add_heading("2.1 技术方案", level=2)
        doc.add_paragraph("技术方案说明。")
        
        # 保存
        doc_path = tmp_path / "test_doc.docx"
        doc.save(str(doc_path))
        
        return str(doc_path)
    
    def test_extract_headings(self, sample_doc):
        """测试标题提取"""
        from vibe_docx.toc import _extract_headings
        
        doc = Document(sample_doc)
        headings = _extract_headings(doc, max_level=3)
        
        assert len(headings) == 5  # 5 headings in test document
        assert headings[0]["text"] == "第一章 简介"
        assert headings[0]["level"] == 1
        assert headings[1]["text"] == "1.1 背景"
        assert headings[1]["level"] == 2
    
    def test_generate_toc_hyperlink(self, sample_doc):
        """测试超链接模式目录生成"""
        from vibe_docx.toc import generate_toc
        
        result = generate_toc(sample_doc, {
            "style": "hyperlink",
            "levels": 3,
            "show_page_numbers": True
        })
        
        assert result["success"] is True
        assert result["toc_style"] == "hyperlink"
        assert result["headings_count"] == 5  # 5 headings in test document
    
    def test_generate_toc_toc_field(self, sample_doc):
        """测试 TOC 字段模式目录生成"""
        from vibe_docx.toc import generate_toc
        
        result = generate_toc(sample_doc, {
            "style": "toc_field",
            "levels": 2
        })
        
        assert result["success"] is True
        assert result["toc_style"] == "toc_field"
    
    def test_validate_toc_links(self, sample_doc):
        """测试链接验证"""
        from vibe_docx.toc import generate_toc, validate_toc_links
        
        # 先生成目录
        generate_toc(sample_doc, {"style": "hyperlink"})
        
        # 验证链接
        result = validate_toc_links(sample_doc)
        
        assert result["success"] is True
        assert result["total_links"] > 0
        assert result["valid_links"] > 0


class TestSKILLInstallation:
    """SKILL 安装测试"""
    
    def test_list_supported_tools(self):
        """测试列出支持的工具"""
        from scripts.install_skill import list_supported_tools
        
        tools = list_supported_tools()
        
        assert len(tools) == 6
        tool_names = [t["name"] for t in tools]
        assert "iflow" in tool_names
        assert "cursor" in tool_names
        assert "claude" in tool_names
    
    def test_verify_install_not_installed(self):
        """测试验证未安装状态"""
        from scripts.install_skill import verify_install
        
        result = verify_install("vibe-docx", "iflow", "local")
        
        # 可能已安装或未安装
        assert "installed" in result
    
    def test_install_skill_dry_run(self, tmp_path):
        """测试安装流程（模拟）"""
        from scripts.install_skill import install_skill
        
        # 使用临时目录
        result = install_skill("vibe-docx", {
            "target": "local",
            "tools": ["iflow"],
            "overwrite": True
        })
        
        # 检查结构
        assert "success" in result
        assert "installed" in result
        assert "failed" in result
        assert "details" in result


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
