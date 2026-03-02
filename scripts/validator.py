# -*- coding: utf-8 -*-
"""
Vibe Docx - Validator Tools

文档分析工具集，用于检测文档问题和结构。
支持两种验证方式：
1. python-docx API 验证（analyze, detect_textboxes 等）
2. XML 直接验证（validate_xml）- 更底层，可检测更多细节
"""

import os
import re
import zipfile
from typing import Dict, List, Any, Optional
from enum import Enum


class IssueSeverity(Enum):
    CRITICAL = "critical"
    WARNING = "warning"
    INFO = "info"


class IssueCategory(Enum):
    FORMAT = "format"
    STRUCTURE = "structure"
    REFERENCE = "reference"
    CONTENT = "content"
    STYLE = "style"


class ValidatorError(Exception):
    """Validator 错误"""
    
    # 统一错误码（DOC 前缀用于文档相关错误）
    ERROR_CODES = {
        # 文档相关 DOC001-DOC010
        "DOC001": {"message": "文档不存在", "recovery": "请提供有效的文档路径"},
        "DOC002": {"message": "不支持的文档格式", "recovery": "仅支持 .docx 格式"},
        "DOC003": {"message": "文档已损坏", "recovery": "尝试修复或使用备份"},
        "DOC004": {"message": "文档被锁定", "recovery": "请关闭其他程序后重试"},
        # 验证相关 VAL001-VAL010
        "VAL001": {"message": "分析失败", "recovery": "检查文件是否损坏"},
        "VAL002": {"message": "XML 解析失败", "recovery": "尝试用 Word 重新保存"},
    }
    
    def __init__(self, code: str, detail: str = ""):
        self.code = code
        self.detail = detail
        info = self.ERROR_CODES.get(code, {"message": "未知错误", "recovery": ""})
        self.message = info["message"]
        self.recovery = info["recovery"]
        super().__init__(f"[{code}] {self.message}: {detail}")


def _check_file(file_path: str) -> None:
    """检查文件有效性"""
    if not os.path.exists(file_path):
        raise ValidatorError("DOC001", file_path)
    if not file_path.lower().endswith('.docx'):
        raise ValidatorError("DOC002", file_path)


def analyze(file_path: str, focus_areas: Optional[List[str]] = None) -> Dict[str, Any]:
    """
    全面分析 DOCX 文档。
    
    Args:
        file_path: DOCX 文件路径
        focus_areas: 可选，关注领域列表 ["format", "structure", "content", "style"]
        
    Returns:
        {
            "success": bool,
            "document_info": {...},
            "issues": [...],
            "stats": {...},
            "risk_factors": [...]
        }
    """
    try:
        _check_file(file_path)
        from docx import Document
        from docx.oxml.ns import qn
        
        doc = Document(file_path)
        paragraphs = list(doc.paragraphs)
        tables = list(doc.tables)
        sections = list(doc.sections)
        
        issues = []
        stats = {
            "paragraphs": len(paragraphs),
            "tables": len(tables),
            "sections": len(sections),
            "images": 0,
            "textboxes": 0,
        }
        
        # === 检测标题样式 ===
        heading_styles: Dict[str, List[int]] = {}
        heading_level_usage: Dict[int, int] = {}  # 统计各级标题使用次数
        
        for i, para in enumerate(paragraphs):
            if para.style.name.startswith('Heading'):
                style_name = para.style.name
                if style_name not in heading_styles:
                    heading_styles[style_name] = []
                heading_styles[style_name].append(i)
                
                # 提取标题级别
                level = 1
                if para.style.name != 'Heading':
                    try:
                        level = int(para.style.name.split()[-1])
                    except:
                        pass
                heading_level_usage[level] = heading_level_usage.get(level, 0) + 1
        
        # 检测标题级别跳跃（如从 H1 直接跳到 H3）
        used_levels = sorted(heading_level_usage.keys())
        if len(used_levels) > 1:
            for i in range(len(used_levels) - 1):
                if used_levels[i + 1] - used_levels[i] > 1:
                    issues.append({
                        "id": "heading_level_skipped",
                        "type": "heading_structure",
                        "category": "structure",
                        "severity": "warning",
                        "detail": f"标题级别跳跃：从 H{used_levels[i]} 跳到 H{used_levels[i+1]}",
                        "auto_fixable": False
                    })
        
        # === 检测表格边框 ===
        tables_without_borders = 0
        for table in tables:
            has_border = False
            try:
                for row in table.rows:
                    for cell in row.cells:
                        borders = cell._element.xpath('.//w:tcBorders')
                        if borders:
                            has_border = True
                            break
                    if has_border:
                        break
            except:
                pass
            if not has_border:
                tables_without_borders += 1
        
        if tables_without_borders > 0:
            issues.append({
                "id": "table_borders_missing",
                "type": "table_no_borders",
                "category": "format",
                "severity": "info",
                "detail": f"发现 {tables_without_borders} 个表格缺少边框",
                "auto_fixable": True
            })
        
        # === 检测空段落 ===
        empty_paragraphs = 0
        for para in paragraphs:
            if not para.text.strip() and len(para.runs) == 0:
                empty_paragraphs += 1
        
        if empty_paragraphs > 3:
            issues.append({
                "id": "excessive_empty_paragraphs",
                "type": "empty_paragraphs",
                "category": "format",
                "severity": "info",
                "detail": f"发现 {empty_paragraphs} 个空段落",
                "auto_fixable": True
            })
        
        # === 检测 Markdown 未转换语法 ===
        markdown_patterns = [
            (r'\*\*[^*]+\*\*', "加粗"),
            (r'\*[^*]+\*', "斜体"),
            (r'~~[^~]+~~', "删除线"),
            (r'`[^`]+`', "代码"),
            (r'#{1,6}\s+\S', "标题"),
        ]
        
        for i, para in enumerate(paragraphs):
            text = para.text
            for pattern, md_type in markdown_patterns:
                import re
                if re.search(pattern, text):
                    issues.append({
                        "id": "markdown_unconverted",
                        "type": "markdown_syntax",
                        "category": "content",
                        "severity": "warning",
                        "detail": f"段落 {i+1} 包含未转换的 Markdown {md_type} 语法",
                        "auto_fixable": True,
                        "location": {"paragraph_index": i}
                    })
                    break  # 每个段落只报告一次
        
        # === 检测图片占位符 ===
        body = doc._body._body
        blips = body.findall('.//' + qn('a:blip'))
        stats["images"] = len(blips)
        
        # 检测文本框
        txbx_contents = body.findall('.//' + qn('w:txbxContent'))
        stats["textboxes"] = len(txbx_contents)
        
        if stats["textboxes"] > 0:
            issues.append({
                "id": "textboxes_detected",
                "type": "textbox",
                "category": "structure",
                "severity": "info",
                "detail": f"检测到 {stats['textboxes']} 个文本框",
                "auto_fixable": True
            })
        
        # === 检测字体不一致 ===
        font_usage: Dict[str, int] = {}
        for para in paragraphs:
            for run in para.runs:
                font_name = run.font.name or "默认"
                font_usage[font_name] = font_usage.get(font_name, 0) + 1
        
        if len(font_usage) > 3:
            top_fonts = sorted(font_usage.items(), key=lambda x: -x[1])[:5]
            issues.append({
                "id": "font_inconsistency",
                "type": "font_mixed",
                "category": "style",
                "severity": "info",
                "detail": f"文档使用了 {len(font_usage)} 种字体: {', '.join([f'{f[0]}({f[1]})' for f in top_fonts])}",
                "auto_fixable": True
            })
        
        # === 检测页面设置 ===
        page_issues = []
        for i, section in enumerate(sections):
            # 检查页边距
            margins = {
                "top": section.top_margin,
                "bottom": section.bottom_margin,
                "left": section.left_margin,
                "right": section.right_margin,
            }
            
            # 检查是否有异常小的页边距（小于 1cm ≈ 360000 EMU）
            from docx.shared import Cm
            min_margin = Cm(1)
            for name, value in margins.items():
                if value and value < min_margin:
                    page_issues.append(f"节 {i+1} {name} 边距过小")
        
        if page_issues:
            issues.append({
                "id": "page_margin_warning",
                "type": "page_setup",
                "category": "format",
                "severity": "warning",
                "detail": "; ".join(page_issues),
                "auto_fixable": True
            })
        
        # === 构建结果 ===
        result = {
            "success": True,
            "document_info": {
                "paragraphs_count": len(paragraphs),
                "tables_count": len(tables),
                "sections_count": len(sections),
                "images_count": stats["images"],
                "textboxes_count": stats["textboxes"],
                "headings_count": sum(heading_level_usage.values()),
                "heading_levels": dict(heading_level_usage),
                "fonts_used": dict(font_usage) if 'font_usage' in dir() else {},
                "file_size": os.path.getsize(file_path)
            },
            "issues": issues,
            "stats": stats,
            "risk_factors": []
        }
        
        # 按关注领域过滤
        if focus_areas:
            result["issues"] = [i for i in result["issues"] if i["category"] in focus_areas]
        
        return result
        
    except ValidatorError as e:
        return {
            "success": False,
            "error": {
                "code": e.code,
                "message": e.message,
                "detail": e.detail,
                "recovery": e.recovery,
                "can_retry": True,
                "error_type": "file_error"
            }
        }
    except Exception as e:
        return {
            "success": False,
            "error": {
                "code": "VAL001",
                "message": "分析失败",
                "detail": str(e),
                "recovery": "检查文件是否损坏，尝试重新保存文档",
                "can_retry": True,
                "error_type": "unknown"
            }
        }


def detect_textboxes(file_path: str) -> Dict[str, Any]:
    """
    检测文档中所有文本框。
    
    Args:
        file_path: DOCX 文件路径
        
    Returns:
        {"success": bool, "textboxes": [...], "stats": {...}}
    """
    try:
        _check_file(file_path)
        from docx import Document
        from docx.oxml.ns import qn
        
        doc = Document(file_path)
        body = doc._body._body
        txbx_contents = body.findall('.//' + qn('w:txbxContent'))
        
        textboxes = []
        has_content_count = 0
        empty_count = 0
        
        for i, txbx in enumerate(txbx_contents):
            paras = txbx.findall('.//' + qn('w:p'))
            para_texts = []
            
            for p in paras:
                runs = p.findall('.//' + qn('w:t'))
                text = ''.join([t.text or '' for t in runs])
                if text.strip():
                    para_texts.append(text.strip())
            
            has_content = len(para_texts) > 0
            if has_content:
                has_content_count += 1
            else:
                empty_count += 1
            
            content_preview = '\n'.join(para_texts)[:200] if para_texts else ""
            
            textboxes.append({
                "index": i,
                "has_content": has_content,
                "paragraph_count": len(para_texts),
                "content_preview": content_preview,
                "paragraphs": para_texts[:5]
            })
        
        return {
            "success": True,
            "textboxes": textboxes,
            "stats": {
                "total_count": len(textboxes),
                "has_content_count": has_content_count,
                "empty_count": empty_count
            }
        }
        
    except ValidatorError as e:
        return {
            "success": False,
            "error": {
                "code": e.code,
                "message": e.message,
                "recovery": e.recovery,
                "can_retry": True,
                "error_type": "file_error"
            }
        }
    except Exception as e:
        return {
            "success": False,
            "error": {
                "code": "VAL003",
                "message": str(e),
                "recovery": "检查文件是否损坏",
                "can_retry": True,
                "error_type": "unknown"
            }
        }


def get_document_structure(file_path: str) -> Dict[str, Any]:
    """
    获取文档结构信息。
    
    Args:
        file_path: DOCX 文件路径
        
    Returns:
        {"success": bool, "structure": {...}}
    """
    try:
        _check_file(file_path)
        from docx import Document
        
        doc = Document(file_path)
        
        structure = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "sections": len(doc.sections),
            "styles_used": list(set(p.style.name for p in doc.paragraphs if p.style))
        }
        
        return {"success": True, "structure": structure}
        
    except Exception as e:
        return {"success": False, "error": str(e)}


def get_section_outline(file_path: str) -> Dict[str, Any]:
    """
    获取章节大纲。
    
    Args:
        file_path: DOCX 文件路径
        
    Returns:
        {"success": bool, "sections": [...]}
    """
    try:
        _check_file(file_path)
        from docx import Document
        
        doc = Document(file_path)
        sections = []
        
        for i, para in enumerate(doc.paragraphs):
            if para.style.name.startswith('Heading'):
                level = 1
                if para.style.name != 'Heading':
                    try:
                        level = int(para.style.name.split()[-1])
                    except:
                        pass
                
                sections.append({
                    "index": len(sections),
                    "level": level,
                    "title": para.text.strip(),
                    "paragraph_index": i
                })
        
        # 检测文本框
        try:
            from docx.oxml.ns import qn
            body = doc._body._body
            txbx_count = len(body.findall('.//' + qn('w:txbxContent')))
            if txbx_count > 0:
                sections.append({
                    "index": len(sections),
                    "level": 0,
                    "title": f"[检测到 {txbx_count} 个文本框]",
                    "paragraph_index": -1
                })
        except:
            pass
        
        return {"success": True, "sections": sections}
        
    except Exception as e:
        return {"success": False, "error": str(e)}


def validate_xml(doc_path: str) -> Dict[str, Any]:
    """
    直接验证 XML 格式（底层验证）。
    
    Args:
        doc_path: DOCX 文件路径
        
    Returns:
        {
            "success": bool,
            "detected": {
                "bold_elements": int,
                "italic_elements": int,
                "underline_elements": int,
                "strikethrough_elements": int,
                "tables": [...],
                "images": [...],
                "headings": [...],
                "page_settings": {...}
            },
            "potential_issues": [...]
        }
    """
    try:
        _check_file(doc_path)
        
        if not zipfile.is_zipfile(doc_path):
            return {"success": False, "error": "不是有效的 DOCX 文件"}
        
        with zipfile.ZipFile(doc_path, 'r') as z:
            try:
                doc_xml = z.read('word/document.xml').decode('utf-8')
            except KeyError:
                return {"success": False, "error": "缺少 word/document.xml"}
            
            rels_xml = ""
            try:
                rels_xml = z.read('word/_rels/document.xml.rels').decode('utf-8')
            except KeyError:
                pass
        
        # 检测格式元素
        detected = {
            "bold_elements": _detect_format_elements(doc_xml, "bold"),
            "italic_elements": _detect_format_elements(doc_xml, "italic"),
            "underline_elements": _detect_format_elements(doc_xml, "underline"),
            "strikethrough_elements": _detect_format_elements(doc_xml, "strikethrough"),
            "tables": _detect_tables(doc_xml),
            "images": _detect_images(doc_xml, rels_xml),
            "headings": _detect_headings(doc_xml),
            "page_settings": _detect_page_settings(doc_xml)
        }
        
        # 查找潜在问题
        potential_issues = _find_potential_issues(doc_xml, detected)
        
        return {
            "success": True,
            "detected": detected,
            "potential_issues": potential_issues
        }
        
    except ValidatorError as e:
        return {
            "success": False,
            "error": {
                "code": e.code,
                "message": e.message,
                "recovery": e.recovery,
                "can_retry": True,
                "error_type": "file_error"
            }
        }
    except Exception as e:
        return {
            "success": False,
            "error": {
                "code": "VAL003",
                "message": "XML 解析失败",
                "detail": str(e),
                "recovery": "检查文件是否损坏，尝试用 Word 重新保存",
                "can_retry": True,
                "error_type": "parse_error"
            }
        }


def _detect_format_elements(doc_xml: str, format_type: str) -> int:
    """检测指定格式的元素数量"""
    tag_map = {
        "bold": r"<w:b\b[^/]*/?>",
        "italic": r"<w:i\b[^/]*/?>",
        "underline": r"<w:u\b[^/]*/?>",
        "strikethrough": r"<w:strike\b[^/]*/?>",
    }
    
    pattern = tag_map.get(format_type)
    if not pattern:
        return 0
    
    return len(re.findall(pattern, doc_xml))


def _detect_tables(doc_xml: str) -> List[Dict[str, Any]]:
    """检测表格信息"""
    tables = []
    table_pattern = r"<w:tbl\b[^>]*>(.*?)</w:tbl>"
    
    for i, match in enumerate(re.finditer(table_pattern, doc_xml, re.DOTALL)):
        tbl_xml = match.group(1)
        rows = len(re.findall(r"<w:tr\b[ >]", tbl_xml))
        
        first_row_match = re.search(r"<w:tr\b[^>]*>(.*?)</w:tr>", tbl_xml, re.DOTALL)
        cols = 0
        if first_row_match:
            cols = len(re.findall(r"<w:tc\b[ >]", first_row_match.group(1)))
        
        has_borders = bool(re.search(r"<w:tblBorders>", tbl_xml))
        
        tables.append({
            "index": i,
            "rows": rows,
            "cols": cols,
            "has_borders": has_borders
        })
    
    return tables


def _detect_images(doc_xml: str, rels_xml: str) -> List[Dict[str, Any]]:
    """检测图片信息"""
    images = []
    blip_pattern = r'<a:blip[^>]*r:embed="([^"]+)"'
    
    for match in re.finditer(blip_pattern, doc_xml):
        rid = match.group(1)
        
        extent_match = re.search(
            r'<wp:extent\s+cx="(\d+)"\s+cy="(\d+)"',
            doc_xml[max(0, match.start()-500):match.end()+500]
        )
        
        width = int(extent_match.group(1)) if extent_match else 0
        height = int(extent_match.group(2)) if extent_match else 0
        
        images.append({
            "rid": rid,
            "width": width,
            "height": height
        })
    
    return images


def _detect_headings(doc_xml: str) -> List[Dict[str, Any]]:
    """检测标题信息"""
    headings = []
    para_pattern = r"<w:p\b[^>]*>(.*?)</w:p>"
    
    for match in re.finditer(para_pattern, doc_xml, re.DOTALL):
        para_xml = match.group(1)
        style_match = re.search(r'<w:pStyle\s+w:val="([^"]+)"', para_xml)
        
        if style_match:
            style = style_match.group(1)
            if style.startswith("Heading") or style.startswith("标题"):
                texts = re.findall(r"<w:t\b[^>]*>([^<]*)</w:t>", para_xml)
                text = "".join(texts).strip()
                
                level_match = re.search(r"Heading(\d)|标题(\d)", style)
                level = int(level_match.group(1) or level_match.group(2)) if level_match else 1
                
                headings.append({
                    "level": level,
                    "text": text,
                    "style": style
                })
    
    return headings


def _detect_page_settings(doc_xml: str) -> Dict[str, Any]:
    """检测页面设置"""
    settings = {
        "width": 11906,
        "height": 16838,
        "margins": {"left": 1440, "right": 1440, "top": 1440, "bottom": 1440}
    }
    
    pg_sz_match = re.search(r'<w:pgSz\s+w:w="(\d+)"\s+w:h="(\d+)"', doc_xml)
    if pg_sz_match:
        settings["width"] = int(pg_sz_match.group(1))
        settings["height"] = int(pg_sz_match.group(2))
    
    pg_mar_match = re.search(
        r'<w:pgMar\s+[^>]*w:left="(\d+)"[^>]*w:right="(\d+)"[^>]*w:top="(\d+)"[^>]*w:bottom="(\d+)"',
        doc_xml
    )
    if pg_mar_match:
        settings["margins"] = {
            "left": int(pg_mar_match.group(1)),
            "right": int(pg_mar_match.group(2)),
            "top": int(pg_mar_match.group(3)),
            "bottom": int(pg_mar_match.group(4))
        }
    
    return settings


def _find_potential_issues(doc_xml: str, detected: Dict[str, Any]) -> List[Dict[str, Any]]:
    """查找潜在问题"""
    issues = []
    
    md_patterns = [
        (r"\*\*([^*]+)\*\*", "markdown_bold"),
        (r"\*([^*]+)\*", "markdown_italic"),
        (r"~~([^~]+)~~", "markdown_strikethrough"),
    ]
    
    text_pattern = r"<w:t\b[^>]*>([^<]*)</w:t>"
    for match in re.finditer(text_pattern, doc_xml):
        text = match.group(1)
        for pattern, subtype in md_patterns:
            if re.search(pattern, text):
                issues.append({
                    "type": "markdown_unconverted",
                    "subtype": subtype,
                    "detail": f"可能的 Markdown 语法未转换: {text[:50]}"
                })
                break
    
    for tbl in detected.get("tables", []):
        if not tbl.get("has_borders"):
            issues.append({
                "type": "table_no_borders",
                "detail": f"表格 {tbl['index']} 缺少边框"
            })
    
    placeholder_pattern = r"\{\{image:(rId\d+)\}\}"
    for match in re.finditer(placeholder_pattern, doc_xml):
        rid = match.group(1)
        issues.append({
            "type": "image_placeholder",
            "detail": f"图片占位符未替换: {{{{image:{rid}}}}}"
        })
    
    return issues


__all__ = [
    'ValidatorError',
    'IssueSeverity',
    'IssueCategory',
    'analyze',
    'detect_textboxes',
    'get_document_structure',
    'get_section_outline',
    'validate_xml'
]
